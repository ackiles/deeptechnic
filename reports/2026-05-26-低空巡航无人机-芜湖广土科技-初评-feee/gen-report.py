#!/usr/bin/env python3
"""Generate technical due diligence report for 芜湖广土科技 - as Word docx."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT = "/Users/ackiles/Documents/【6-4】对接项目/安徽节目录制项目/初评报告_低空巡航无人机_芜湖广土科技.docx"

doc = Document()

# ---- Styles ----
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level, sz in [(1, 18), (2, 15), (3, 13)]:
    s = doc.styles[f'Heading {level}']
    s.font.name = 'Arial'
    s.font.size = Pt(sz)
    s.font.bold = True
    s.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
    s.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    s.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    s.paragraph_format.space_after = Pt(10)
    s.paragraph_format.line_spacing = 1.5

def add_para(text, bold=False, color=None, size=None, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_para_runs(runs, space_after=6):
    """runs: list of (text, bold, color, size) tuples"""
    p = doc.add_paragraph()
    for text, bold, color, size in runs:
        run = p.add_run(text)
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if bold:
            run.bold = True
        if color:
            run.font.color.rgb = color
        if size:
            run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.4
    return p

def add_table(headers, rows, col_widths=None):
    """Create a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Blue header background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A3A5C"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    
    # Data rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Arial'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            if ci == 0:
                run.bold = True
            # Alternating row shading
            if ri % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F7FA"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    
    return table

def add_table_title(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    return p

def spacer():
    return doc.add_paragraph()

# ========== TITLE PAGE ==========
for _ in range(6):
    doc.add_paragraph()

add_para('Deeptechnic 技研', bold=True, color=RGBColor(0x1A, 0x3A, 0x5C), size=26, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('技术初评报告', color=RGBColor(0x66, 0x66, 0x66), size=20, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
add_para('芜湖广土科技有限责任公司', bold=True, color=RGBColor(0x1A, 0x3A, 0x5C), size=24, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
add_para('低空巡航无人机及配套产品线技术研判', color=RGBColor(0x44, 0x44, 0x44), size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

# Separator line
add_para('─' * 40, color=RGBColor(0xCC, 0xCC, 0xCC), size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
add_para('报告编号：DT-2026-001', color=RGBColor(0x66, 0x66, 0x66), size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('报告日期：2026年5月26日', color=RGBColor(0x66, 0x66, 0x66), size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('密级：机密 — 仅限内部传阅', color=RGBColor(0xCC, 0x33, 0x33), size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

doc.add_page_break()

# ========== TOC ==========
doc.add_heading('目录', level=1)
toc_items = [
    '一、项目概述',
    '  1.1 公司基本情况',
    '  1.2 核心产品线概览',
    '二、核心产品线与技术评估',
    '  2.1 RemoteID 收发一体模块',
    '  2.2 4+1 固定翼巡航无人机平台',
    '  2.3 定制产品线评估',
    '三、产业化背景与政策环境',
    '四、市场竞争格局',
    '五、团队评估',
    '六、风险识别',
    '七、综合研判与结论',
    '八、后续尽调建议',
]
for item in toc_items:
    add_para(item, space_after=2)

doc.add_page_break()

# ========== SECTION 1 ==========
doc.add_heading('一、项目概述', level=1)
add_para(
    '本报告对芜湖广土科技有限责任公司（以下简称「广土科技」）的低空巡航无人机项目进行独立第三方技术初评。'
    '评估基于项目商业计划书（BP）内容及公开可查信息，依据 Deeptechnic（技研）技术研判框架，'
    '从技术先进性、产品完整性、市场前景、团队胜任力和风险维度展开系统分析。'
)

doc.add_heading('1.1 公司基本情况', level=2)

add_table(
    ['项目', '内容'],
    [
        ['公司名称', '芜湖广土科技有限责任公司'],
        ['公司状态', '一人公司（单一工程师主导）'],
        ['联系地址', '芜湖航空产业园'],
        ['联系人', '庄金峰'],
        ['邮箱', 'jinfeng.Zhuang@aliyun.com'],
        ['手机', '15021315961'],
        ['融资需求', '100万元（12-15个月）'],
        ['盈利模式', '通过技术服务而非生产制造盈利'],
    ],
    col_widths=[4, 12]
)
add_table_title('表1：公司基本概况')

doc.add_heading('1.2 核心产品线概览', level=2)
add_para('广土科技目前规划了六条产品线，呈现「一条主线+多条辅线」的结构：')
add_bullet('RemoteID 收发一体模块：首款产品，响应民航局远程识别强制要求')
add_bullet('4+1 固定翼巡航无人机平台：核心产品，3D打印，翼展2m，120km航程')
add_bullet('NPU芯片（联合研发）：提供软件方案+两个应用场景（无人机+AI硬件）')
add_bullet('MicroLED光源芯片：汽车ADB大灯控制算法（与上海公司合作）')
add_bullet('AI智能硬件：语音玩具+大模型对话+定制语音')
add_bullet('定制化物联网模块：空压机振动监测等项目式定制')

add_para_runs([
    ('综合判断：', True, None, 11),
    ('[中-差]', True, RGBColor(0xB8, 0x5C, 0x00), 11),
    ('  产品方向定位符合行业趋势，但六条线并行远超一人公司承载能力。', False, None, 10),
])

doc.add_page_break()

# ========== SECTION 2 ==========
doc.add_heading('二、核心产品线与技术评估', level=1)

# 2.1 RemoteID
doc.add_heading('2.1 RemoteID 收发一体模块', level=2)

doc.add_heading('方案概述', level=3)
add_para(
    'RemoteID模块被BP定位为首款产品，面向中国民航局无人机远程识别强制标准。'
    '模块集成4G上云、GPS定位、气压计、WiFi等通信与感知功能，'
    '支持机载端实时广播和地面站接收解析的双向能力。'
)

doc.add_heading('技术方案评估', level=3)
add_para(
    '方案采用「WiFi近距离广播+4G远程上云」的混合路径，这是当前Remote ID模块的主流技术路线。'
    'GPS+气压计双模定位提供了冗余设计。方案总体技术思路合理。'
)

doc.add_heading('竞品对比', level=3)
add_table(
    ['参数', '广土科技', 'Dronetag', 'uAvionix'],
    [
        ['通信方式', 'WiFi+4G', 'BLE+WiFi', 'WiFi'],
        ['定位方式', 'GPS+气压计', 'GPS+气压计', 'GPS'],
        ['价格区间', '未披露', '€99-199', '$199-349'],
        ['认证状态', '未认证', 'FAA/欧盟', 'FAA/欧盟'],
        ['量产状态', '未量产', '已量产', '已量产'],
    ],
    col_widths=[3.5, 4.5, 4, 4]
)
add_table_title('表2：Remote ID模块竞品对比')

doc.add_heading('风险评估', level=3)
add_bullet('市场窗口期有限：DJI等主流无人机已全系预装Remote ID，独立模块市场需求集中在存量无人机改造，预计2026年后趋于饱和')
add_bullet('认证成本：民航局合规认证流程耗时耗钱（预计数万至十余万元），100万元融资中可能被大量占用')
add_bullet('竞争压价：Remote ID独立模块领域已有Dronetag、Holybro等国际品牌，国内珠三角代工厂很快会跟进低成本方案')

add_para_runs([
    ('技术评级：', True, None, 11),
    ('中  ', True, RGBColor(0xD4, 0x73, 0x0D), 11),
    ('方案合理，但差异化不足，市场窗口有限。', False, None, 10),
])

spacer()

# 2.2 Drone Platform
doc.add_heading('2.2 4+1 固定翼巡航无人机平台', level=2)

doc.add_heading('方案概述', level=3)
add_para(
    '广土科技的核心产品：3D打印的4+1复合翼（垂直起降固定翼）无人机，翼展2m，'
    '宣称航程120km、续航2小时。计划搭载可变焦距+NPU摄像头、红外相机和4G图传模块，'
    '用于高速/大农田/水域/边境/森林/海岸线巡检。'
)

doc.add_heading('技术方案评估', level=3)
add_para(
    '构型选择（4+1 VTOL复合翼）合理，是当前10-50kg级工业无人机的标准构型。'
    '3D打印机身降低了起步门槛，但带来了量产一致性和结构耐久性的疑问。'
)

doc.add_heading('竞品对比', level=3)
add_table(
    ['指标', '广土科技', '纵横CW-15', '科比特插翅虎', 'DJI M350'],
    [
        ['类型', '4+1复合翼', '纯固定翼', '复合翼', '多旋翼'],
        ['翼展', '2m', '3.4m', '2.5m', '-'],
        ['航程', '120km', '200km', '180km', '20km'],
        ['续航', '2h', '4h', '3.5h', '55min'],
        ['载荷', '未披露', '2kg', '2kg', '2.7kg'],
        ['售价', '未披露', '~15万', '~12万', '~6万'],
        ['量产', '未量产', '已量产', '已量产', '已量产'],
    ],
    col_widths=[2.8, 3.2, 3.2, 3.2, 3.2]
)
add_table_title('表3：固定翼巡航无人机竞品对比')

doc.add_heading('关键风险评估', level=3)
add_bullet('核心参数未经第三方实测验证——BP中未提供试飞视频或测试报告')
add_bullet('3D打印结构在持续振动环境下的可靠性在行业中尚无大规模验证')
add_bullet('飞控系统方案未披露（自研/Pixhawk/ArduPilot），这是衡量底层技术能力的关键')
add_bullet('载重能力未明确，严重制约了载荷配置空间')

add_para_runs([
    ('技术评级：', True, None, 11),
    ('[中-差]  ', True, RGBColor(0xB8, 0x5C, 0x00), 11),
    ('构型选择正确但参数偏下，缺实测数据，量产路径不清晰。', False, None, 10),
])

spacer()

# 2.3 Custom Products
doc.add_heading('2.3 定制产品线评估', level=2)

doc.add_heading('NPU芯片联合研发', level=3)
add_para(
    'BP中描述与第三方联合研发NPU芯片，广土科技提供软件方案和两个应用场景（无人机+AI智能硬件）。'
    '此模式在行业常见，但一人公司承担NPU全套软件栈开发（编译器、驱动、算子库）可信度存疑——'
    '该工作通常需要5-10人团队。合作方身份、合作阶段、IP分配方案均未披露。'
)

doc.add_heading('MicroLED光源芯片', level=3)
add_para(
    '与上海公司合作开发汽车大灯投影方案的底层控制算法。方向与技术趋势一致（ADB成为中高端车型标配），'
    '但车规级软件需满足ISO 26262功能安全标准，一人公司的软件质量保证能力堪忧。'
)

doc.add_heading('AI智能硬件（语音玩具）', level=3)
add_para(
    'AI语音玩具赛道2024-2025年极度拥挤（字节、百度、科大讯飞等大厂均已入局）。'
    '技术门槛低（大模型API调用），无需硬件突破。作为被动定制接单项目可以理解，但作为主动产品线没有竞争力。'
)

doc.add_heading('IoT物联网模块', level=3)
add_para('BP中明确表示采用成熟模块部署，不自研。这一态度务实合理。')

add_para_runs([
    ('综合评级：', True, None, 11),
    ('[差]  ', True, RGBColor(0xC0, 0x39, 0x2B), 11),
    ('产品线高度分散，核心方向被稀释。', False, None, 10),
])

doc.add_page_break()

# ========== SECTION 3 ==========
doc.add_heading('三、产业化背景与政策环境', level=1)

doc.add_heading('3.1 低空经济政策大势', level=2)
add_para(
    '低空经济已上升为国家战略。2023年中央经济工作会议首提、2024年写入政府工作报告、'
    '2025年《无人驾驶航空器飞行管理暂行条例》全面实施。'
    '民航局预测2030年低空经济对GDP贡献超2万亿元。'
    '安徽芜湖被列为低空经济试点城市，芜湖航空产业园是安徽省重点产业集聚区。'
)

doc.add_heading('3.2 对广土科技的影响', level=2)
add_bullet('政策大方向利好，行业处于高速上升期')
add_bullet('芜湖航空产业园有租金减免、人才补贴等扶持政策（实际落实程度有待创始人确认）')
add_bullet('低空经济投资热带动了资本关注，但资本更偏好已量产和有规模验证的企业')

add_para_runs([
    ('机遇评级：', True, None, 11),
    ('[良]  ', True, RGBColor(0x2B, 0x57, 0x9A), 11),
    ('赛道正确，但窗口期不会永远敞开（2025-2028是关键窗口期）。', False, None, 10),
])

# ========== SECTION 4 ==========
doc.add_heading('四、市场竞争格局', level=1)

doc.add_heading('4.1 竞争层级分析', level=2)
add_bullet('第一梯队（A级）：大疆创新（全能霸主）、纵横股份（复合翼巡检龙头，上市企业）')
add_bullet('第二梯队（B级）：科比特航空、云圣智能（已获数亿融资）')
add_bullet('第三梯队（C级）：普宙飞行器、各区域型小团队')
add_bullet('广土科技潜在定位：C级以下，尚未进入竞争棋盘')

doc.add_heading('4.2 差异化空间', level=2)
add_para('广土科技可能的差异化路径：')
add_bullet('极致成本：3D打印+精简团队，整机价格可能显著低于竞品')
add_bullet('本地化服务：芜湖航空产业园贴近长三角客户，现场响应速度快')
add_bullet('灵活定制：小批量、快速改型的能力（「2周打板+2周调试」）')

add_para_runs([
    ('竞争评级：', True, None, 11),
    ('[差]  ', True, RGBColor(0xC0, 0x39, 0x2B), 11),
    ('头部玩家已建立品牌/渠道/资金壁垒，差异化路径需要极强的执行力。', False, None, 10),
])

doc.add_page_break()

# ========== SECTION 5 ==========
doc.add_heading('五、团队评估', level=1)

doc.add_heading('5.1 创始人能力分析', level=2)
add_bullet('从BP技术描述判断：有较完整的嵌入式系统开发经验（GPS/4G/WiFi/传感器/3D打印）')
add_bullet('从BP的融资方案和五年规划描述判断：商业逻辑和产品思维有一定的成熟度')
add_bullet('从个人邮箱(aliyun.com)判断：尚未建立企业级IT基础设施')

doc.add_heading('5.2 一人公司模式的根本性挑战', level=2)
add_table(
    ['挑战维度', '具体描述', '严重度'],
    [
        ['技术单点', '关键技术决策无交叉验证，错误难以被及时发现', '高'],
        ['精力分散', '六条产品线同时推进，单人每周168h也不够', '极高'],
        ['供应链', 'BOM采购、备货、品控全部亲力亲为', '高'],
        ['适航认证', '民航认证流程极为繁琐，单人难以应对', '高'],
        ['商务对接', '同时负责技术和商务，研发效率下行', '中'],
        ['团队替代', '创始人不可替代本身就是投资风险', '中'],
    ],
    col_widths=[3, 9, 3]
)
add_table_title('表4：一人公司模式风险清单')

add_para_runs([
    ('团队评级：', True, None, 11),
    ('[中-差]  ', True, RGBColor(0xB8, 0x5C, 0x00), 11),
    ('创始人展现出一定的技术基础，但一人模式在硬件赛道风险极高。', False, None, 10),
])

doc.add_page_break()

# ========== SECTION 6 ==========
doc.add_heading('六、风险识别', level=1)

doc.add_heading('6.1 技术风险', level=2)
add_bullet('【高】产品从未量产，技术债全部由一人承担')
add_bullet('【高】无人机的核心参数（120km/2h）是否实测达到？缺少第三方验证')
add_bullet('【中】3D打印飞机的结构耐久性和量产一致性有待验证')
add_bullet('【中】飞控系统方案不明确（自研/开源/商用？）')

doc.add_heading('6.2 市场风险', level=2)
add_bullet('【高】Remote ID模块市场窗口期有限，预计2026年后饱和')
add_bullet('【高】工业无人机巡检市场已被大疆/纵横/科比特等先发者占据')
add_bullet('【中】AI语音玩具属于消费电子红海，一人公司无竞争力')
add_bullet('【中】NPU芯片联合研发的场景出货量（50-100套）与芯片流片量级严重不匹配')

doc.add_heading('6.3 团队风险', level=2)
add_bullet('【高】关键技术全部依赖创始人一人，健康/家庭/个人风险传导至公司')
add_bullet('【高】20万元人力预算不足以支撑12-15个月的工程师薪酬')
add_bullet('【中】无法核实创始人的详细技术履历')

doc.add_heading('6.4 财务与融资风险', level=2)
add_bullet('【高】100万元融资在12-15个月内需覆盖认证、打样、试飞、人力、市场等全部开销，资金缺口明显')
add_bullet('【中】初创企业以个人邮箱联系客户，会降低合作伙伴的信任度')

doc.add_page_break()

# ========== SECTION 7 ==========
doc.add_heading('七、综合研判与结论', level=1)

doc.add_heading('7.1 综合评级', level=2)
add_table(
    ['评估维度', '评级', '简要说明'],
    [
        ['技术先进性', '[中-差]', '跟随策略，无突破性创新'],
        ['技术成熟度', '[中-差]', '无量产经验，Remote ID未出货'],
        ['产品系统完整性', '[中]', '方向思路合理，但六条线分散'],
        ['市场竞争力', '[差]', '头部玩家已建立产业壁垒'],
        ['团队胜任力', '[中-差]', '创始人个人能力可，但一人模式风险极高'],
        ['商业化前景', '[差]', '直销+本地服务模式天花板明显'],
    ],
    col_widths=[3.5, 3, 9]
)
add_table_title('表5：综合评级矩阵')

spacer()
add_para_runs([
    ('综合评级：', True, RGBColor(0x1A, 0x3A, 0x5C), 16),
    ('  ', False, None, 4),
    ('[中-差]', True, RGBColor(0xB8, 0x5C, 0x00), 16),
], space_after=12)

add_para(
    '广土科技的项目处于极早期的种子前阶段。技术方向选择正确（Remote ID + 固定翼巡航无人机 + AI边缘计算顺应行业趋势），'
    '创始人展现出一定的嵌入式系统工程能力。但一人公司的资源极限、六条并行产品线的战略失焦、'
    '无任何量产和认证记录的核心短板、以及工业无人机赛道的激烈竞争格局，使得该项目在当前状态下'
    '技术风险显著高于可接受水平。'
)

doc.add_heading('7.2 结论倾向', level=2)
add_para('建议保持关注，但不建议在当前阶段直接投资。可在以下里程碑达成后重新评估：')
add_bullet('Remote ID模块完成民航合规认证并通过第三方测试')
add_bullet('固定翼巡航无人机完成至少50小时无故障试飞')
add_bullet('创始人完成至少1-2名关键技术人员招聘')
add_bullet('聚焦为核心2-3条产品线，而非现在的6条')

doc.add_page_break()

# ========== SECTION 8 ==========
doc.add_heading('八、后续尽调建议', level=1)

doc.add_heading('8.1 建议尽调事项', level=2)

doc.add_heading('技术尽调', level=3)
add_bullet('【必查】创始人软件/嵌入式代码能力：要求提供GitHub/Coding等平台的代码库')
add_bullet('【必查】Remote ID模块的原理样机演示和射频测试报告')
add_bullet('【必查】无人机飞控系统的选型和底层代码来源')
add_bullet('【必查】NPU芯片联合研发的合作协议和IP归属条款')
add_bullet('【可选】MicroLED控制算法的演示和功能安全文档')

doc.add_heading('商业尽调', level=3)
add_bullet('【必查】芜湖广土科技的工商注册信息（注册资本、实缴、股东结构）')
add_bullet('【必查】创始人的学历证明和过往工作履历')
add_bullet('【必查】NPU和MicroLED合作方的身份和资质')
add_bullet('【必查】已有的客户意向书或合作备忘录（MOU）')
add_bullet('【可选】100万元融资的详细资金运用计划')

doc.add_heading('法律尽调', level=3)
add_bullet('【必查】专利/软著申请清单（BP中未提知识产权布局）')
add_bullet('【必查】技术来源的合规性（是否涉及前雇主的知识产权？）')
add_bullet('【必查】联合研发协议中已签署的合作文件')

doc.add_heading('8.2 建议后续跟踪节点', level=2)
add_bullet('Remote ID模块通过民航认证测试（预计需3-6个月）')
add_bullet('完成至少1次公开的固定翼无人机试飞演示')
add_bullet('招聘到第1名嵌入式全职工程师')
add_bullet('获取首个Remote ID模块采购订单')

# Closing
spacer()
spacer()
add_para('— 报告完 —', color=RGBColor(0x99, 0x99, 0x99), size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
add_para('本报告由Deeptechnic（技研）基于BP文件及公开可查信息编制', color=RGBColor(0x99, 0x99, 0x99), size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para('所有判断均为技术研判意见，不构成投资建议', color=RGBColor(0x99, 0x99, 0x99), size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

# Save
doc.save(OUTPUT)
print(f"Report saved: {OUTPUT}")
