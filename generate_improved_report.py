#!/usr/bin/env python3
"""Generate improved 技术初评报告 for 普瑞达 Primdata AI (AI多智能体科研助理)"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ======== STYLES ========
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(text, bold=False, italic=False, size=None, color=None, alignment=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = size
    if color: run.font.color.rgb = color
    if alignment: p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = Pt(18)
    return p

def add_styled_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = '微软雅黑'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        set_cell_shading(cell, '2F5496')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255,255,255)
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx+1].cells[c_idx]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = '微软雅黑'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    # Set col widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    return table

# ======== COVER ========
for _ in range(6):
    doc.add_paragraph()

add_para('技术初评报告', bold=True, size=Pt(26), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
add_para('Deeptechnic（技研）', bold=True, size=Pt(16), alignment=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(47, 84, 150), space_after=Pt(8))
add_para('技术尽职调查 · 独立第三方技术研判', size=Pt(11), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(30))

add_para('项目名称：AI多智能体科研助理', bold=True, size=Pt(12), alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('公司名称：合肥普瑞达智能科技有限公司（Primdata AI）', size=Pt(11), alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('报告日期：2026年5月27日（更新版）', size=Pt(11), alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('密级：商业机密 · 仅限指定对象阅读', size=Pt(10), alignment=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(192, 0, 0))

doc.add_page_break()

# ======== 报告概述页 ========
add_para('', size=Pt(6))
add_para('报告概述', bold=True, size=Pt(18), alignment=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(47, 84, 150), space_after=Pt(16))

add_para('一、项目画像', bold=True, size=Pt(12), color=RGBColor(47, 84, 150), space_after=Pt(6))
add_para('合肥普瑞达智能科技（Primdata AI）——「生命科学科研人员的AI数字员工」。种子轮/天使轮，2人团队（中科大博士后+生信专家），产品已内测。聚焦生信数据分析、文献调研、科研辅助三大场景，原生嵌入飞书/微信/企微。')

add_para('二、核心问题', bold=True, size=Pt(12), color=RGBColor(47, 84, 150), space_after=Pt(6))
add_para('① 技术壁垒是否足够？应用层Agent方案能否在市场立足？')
add_para('② 2人科研背景团队能否同时驱动产品研发和商业化？')
add_para('③ 在Elicit/Scite/PaperQA/赛舵智能等竞品环伺下，差异化优势是否可持续？')
add_para('④ 机构License的6-12个月采购周期对种子轮项目意味着怎样的现金流风险？')
add_para('⑤ 2026年AI Agent产业高速演进背景下，普瑞达如何避免被通用Agent平台吞噬？')

add_para('三、核心结论', bold=True, size=Pt(12), color=RGBColor(47, 84, 150), space_after=Pt(6))

add_styled_table(
    ['维度', '评级', '一句话判断'],
    [
        ['技术先进性', '★★★☆☆', '应用层Agent合理但壁垒不足，RAG管线质量是成败关键'],
        ['团队能力', '★★★☆☆', '科研背景贴合痛点，但缺少工程化和商业化角色是最大短板'],
        ['商业化前景', '★★★☆☆', '个人付费已跑通（亮点），但机构采购周期长是核心风险'],
        ['竞品地位', '★★★☆☆', '定位精准但在竞品中无明显护城河，需靠执行速度取胜'],
        ['学术基础', '★★★☆☆', '生信Agent方向学术界已有多项对标工作，技术可行但无先发优势'],
    ],
    col_widths=[2.5, 2, 12]
)

add_para('')
add_para('───────────────────────────────────────────', size=Pt(8))
add_para('总体评级：★★★☆☆（中等偏上）——关注级项目，建议推进尽调但重点关注团队补齐和商业化验证。', bold=True, size=Pt(11), alignment=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(47, 84, 150))
add_para('')

add_para('四、关键风险清单（按严重程度排序）', bold=True, size=Pt(12), color=RGBColor(192, 0, 0), space_after=Pt(6))

add_styled_table(
    ['排序', '风险', '等级', '简要说明'],
    [
        ['1', '关键人员依赖', '🔴 高', '仅2人团队，任何一位核心成员退出将严重影响项目推进'],
        ['2', '机构采购周期长', '🔴 高', '机构License转化周期6-12个月，种子轮现金流可能断裂'],
        ['3', '技术壁垒薄弱', '🟡 中高', '应用层Agent方案门槛低，通用Agent平台快速进化可能侵蚀市场'],
        ['4', 'LLM幻觉风险', '🟡 中高', '科研文献引用错误可能导致严重后果，RAG管线可靠性尚未验证'],
        ['5', '竞品压力', '🟡 中', '国际+国内多款竞品已获融资和用户，普瑞达差异化优势待验证'],
        ['6', '商业化能力', '🟡 中高', '纯科研背景团队缺少商业运营经验'],
    ],
    col_widths=[1, 4, 1.5, 10]
)

doc.add_page_break()

# ======== 1. 总体评价 ========
add_heading_styled('一、总体评价与建议', level=1)

add_para('总体评分：', bold=True, size=Pt(14))
# Rating table
add_styled_table(
    ['技术先进性', '团队能力', '商业化前景', '整体评分'],
    [['★★★☆☆', '★★★☆☆', '★★★☆☆', '★★★☆☆']],
    col_widths=[3.5, 3.5, 3.5, 3.5]
)
add_para('评级说明：★★★☆☆（中等偏上）', bold=True, color=RGBColor(47, 84, 150))

add_para('')
add_para('普瑞达（Primdata AI）定位为「生命科学科研人员的AI数字员工」，聚焦AI Agent在生物信息学分析、文献调研、科研辅助等场景的应用。本项目处于种子轮/天使轮阶段，2人创始团队，产品已内测。本次更新版本增强了产业动态分析、学术前沿对比和竞品深度解读。')

# Highlights
add_heading_styled('核心亮点', level=2)
highlights = [
    '痛点精准：生命科学领域研究者（尤其是生物信息学方向）面临文献过载、分析工具门槛高、工作流碎片化等真实痛点，普瑞达的「AI数字员工」定位切中实际需求。',
    '团队背景匹配：两位创始人分别具备中国科学技术大学博士后科研背景和生物信息学分析专长，「我们是用户的原型」这一表述表明团队对痛点有切身理解。',
    '原生嵌入工作流：原生接入飞书/微信/企业微信，降低了科研人员的使用门槛，这一产品设计思路正确。',
    '市场信号积极：核心用户使用频次高、学生付费已跑通、自媒体获客成本低于10元/人——这些信号值得肯定。',
    '产业趋势支撑：2026年5月，AI Agent赛道正处于高速发展期，国产Agent模型（如昆仑万维SkyClaw-v1.0）进入全球第一梯队，为应用层创业提供了更低成本的模型基础设施。',
]
for h in highlights:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(h)
    run.font.size = Pt(10.5)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# Concerns
add_heading_styled('核心关注点', level=2)
concerns = [
    '技术壁垒不足：当前产品构建在基础大模型API之上做应用层Agent，技术栈属于行业通行做法，难以形成持久技术壁垒。学术界已有多个对标工作（AutoBA 2023、BRAD 2024、oxo-call 2026等），产业端同类产品众多。',
    '团队规模过小：仅2人团队，且均为科研背景，缺少AI工程化/AI Agent系统架构方向的专职工程师，是当前最突出的结构性短板。',
    '竞品压力明确：国际上有Elicit/Scite/PaperQA/Consensus等产品，国内有赛舵智能等竞品，DeekSeek等团队也在探索科研自动化方向。差异化优势尚待验证。',
    '商业化路径依赖机构采购：核心收入来源为机构License，而机构采购周期长（6-12个月），可能影响现金流。学生付费虽已跑通，但客单价低。',
]
for c in concerns:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(c)
    run.font.size = Pt(10.5)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ======== 2. 公司概况 ========
add_heading_styled('二、公司概况与项目定位', level=1)
add_heading_styled('2.1 基本信息', level=2)

add_styled_table(
    ['项目', '内容'],
    [
        ['公司全称', '合肥普瑞达智能科技有限公司'],
        ['品牌名称', 'Primdata AI'],
        ['产品定位', '生命科学科研人员的AI数字员工\n（文献调研·数据分析·研究辅助）'],
        ['技术路线', '应用层Agent（不做基础模型）'],
        ['核心产品形态', '原生嵌入飞书/微信/企业微信'],
        ['团队规模', '2人（中科大博士后+生物信息学专家）'],
        ['融资阶段', '种子轮/天使轮'],
        ['核心壁垒主张', 'AI自动核实每一篇文献的审查机制'],
    ],
    col_widths=[3.5, 12]
)

add_heading_styled('2.2 项目定位分析', level=2)
add_para('普瑞达明确表示「不做基础模型，做应用层Agent」，定位清晰、务实。这一选择合理——在当前大模型生态下，应用层创业公司资源有限，应专注于垂直场景的深度适配。公司聚焦「生命科学」这一垂直领域，避免了与通用AI助手的正面竞争。')
add_para('其「原生嵌入飞书/微信/企业微信」的产品形态，降低了科研人员的认知门槛——用户无需离开日常通讯工具即可完成文献调研、数据分析等任务。这一设计体现了对目标用户使用习惯的深入理解。')
add_para('深入分析：', bold=True)
add_para('• 产品覆盖的四个具体子场景（生信数据分析、文献调研追踪、生信教学/入门、临研辅助）均为高频率真实需求，体现了团队对科研工作流的具体理解。')
add_para('• 「不做基础模型」的战略选择意味着普瑞达的核心竞争力将来自产品体验、领域知识库和用户粘性，而非底层技术能力。这要求团队在产品迭代速度上保持领先。')

doc.add_page_break()

# ======== 3. 技术维度深度分析 ========
add_heading_styled('三、技术维度深度分析', level=1)

add_heading_styled('3.1 核心技术架构评估', level=2)
add_para('普瑞达的技术方案属于「大模型应用层Agent」架构，由以下核心模块组成：')

add_para('(1) 大模型接入层', bold=True)
add_para('对接GPT-4/Claude/文心一言等基础大模型API。这是当前AI应用创业的通行做法，技术门槛不高。')
add_para('【2026年行业更新】2026年5月，昆仑万维发布SkyClaw-v1.0国产Agent模型，百万token上下文，深度适配Agent工作场景。这意味着普瑞达可依赖更低成本的国产模型。行业经验显示，从GPT-4o切换到DeepSeek等国产模型可降低97%的API成本（Dev.to社区报告，2026年5月）。')

add_para('(2) Agent编排引擎', bold=True)
add_para('设计多步骤工作流的编排机制。行业已有LangChain/LlamaIndex/AutoGen等成熟框架可用。2026年Agent开发范式趋于成熟——Skill(稳定执行)→SubAgent(调度控制)→Agent Team(协作提效)三层架构成为行业共识（开源中国，2026年5月）。行业挑战在于多Agent协作时的状态管理和一致性控制。')

add_para('(3) 文献审查机制（核心差异化）', bold=True)
add_para('BP强调的「AI自动核实每一篇文献」是差异化亮点，但技术实现难度高——需要构建高质量的检索增强生成（RAG）管线，包含知识库构建、引用溯源、幻觉检测等环节。科研文献的准确性要求极高——LLM幻觉可能导致严重后果（如误导实验设计）。')
add_para('学术界前沿：oxo-call（2026年4月）展示了通过文档增强的RAG方法大幅提升LLM在生物信息学领域的命令行生成准确性，验证了RAG技术路线的有效性。')

add_para('(4) 本地化部署能力', bold=True)
add_para('面向高合规客户（医院/药企）提供数据不出域的私有化部署方案。2026年容器化技术高度成熟，技术可行但需考虑部署后的维护成本。')

add_para('(5) 飞书/微信集成', bold=True)
add_para('原生嵌入IM工具，涉及开放平台API集成，技术难度中等。')

add_heading_styled('3.2 技术成熟度判断', level=2)
add_styled_table(
    ['技术模块', '成熟度', '说明'],
    [
        ['文献调研摘要', '★★★☆', 'RAG+LLM摘要，行业有成熟方案，但科研文献的准确性要求极高，RAG管线质量是关键瓶颈'],
        ['生信数据分析', '★★☆☆', 'AI自动执行生信分析pipeline是前沿方向（如AutoBA, oxo-call等学术工作），通用化困难，需要深厚的领域知识库支撑'],
        ['科研辅助/写作', '★★★☆', 'LLM辅助写作相对成熟，但DeepSeek的自动研究Skill（2026年5月）显示该领域正在快速进化'],
        ['Agent编排引擎', '★★☆☆', '使用行业框架快速搭建可行，但生产环境稳定性和可扩展性需要大量工程打磨'],
        ['本地化部署', '★★★☆', '容器化方案成熟，但小团队的维护成本高'],
        ['IM集成', '★★★★', '飞书/微信API集成技术成熟'],
    ],
    col_widths=[3.5, 2, 10.5]
)

add_heading_styled('3.3 算法与模型选择合理性', level=2)
add_para('BP中未明确披露所使用的基础模型和Agent框架。从行业最佳实践来看：')
add_para('• 对于文献分析/生成类任务，GPT-4/Claude 3.5在英文文献处理上有优势，但中文文献可能需要国产模型（如DeepSeek、GLM-4）补充。')
add_para('• 对于生信代码生成任务，当前LLM在生物信息学工具的命令行生成方面已有突破（oxo-call, 2026年4月），技术可行性得到学术验证。')
add_para('• Agent框架选择方面，建议关注AutoBA（2023年）——一个专门面向生物信息学分析的自主AI Agent系统。')
add_para('• 2026年趋势：国产Agent模型快速成熟（昆仑万维SkyClaw-v1.0），模型成本持续降低，对应用层创业者是利好。')

doc.add_page_break()

# ======== 4. 学术前沿 ========
add_heading_styled('四、学术前沿与对标研究', level=1)
add_para('通过学术文献检索，梳理了以下与本项目高度相关的学术前沿工作：')

add_styled_table(
    ['论文/工作', '时间', '相关性说明', '核心启示'],
    [
        ['AutoBA\n(Auto Bioinformatics Analysis)', '2023.09', '首个面向生信分析的自主AI Agent，与普瑞达的生信分析场景高度重合', '学术界证明LLM驱动的生信分析Agent技术可行'],
        ['BRAD\n(LLM Digital Biology)', '2024.09', '语言模型驱动的数字生物学平台，展示LLM在生命科学领域的应用潜力', 'LLM+生信工具链集成是研究热点'],
        ['LLMs in Bioinformatics:\nA Survey', '2025.03', 'LLM在生物信息学领域的全面综述', 'LLM在序列分析、结构预测等领域表现优异，但端到端工作流仍有大量开放问题'],
        ['oxo-call', '2026.04', '文档增强的生信命令行生成，通过RAG提升LLM对生信工具的理解', 'RAG可大幅提升LLM在生信场景的准确性'],
        ['Agentic AI for\nNGS Analysis', '2025.12', '面向NGS下游分析的AI Agent模型，直接竞品级学术工作', '生信数据分析场景已被学术界前沿团队关注'],
        ['SciAtlas', '2026.05', '大规模科研知识图谱用于自动化科学研究', '科研自动化方向学术工作加速'],
        ['AutoResearchBench', '2026.04', '科研Agent标准化评测基准', '科研Agent方向已有标准化评估方法'],
    ],
    col_widths=[3.5, 1.5, 5, 5.5]
)

add_para('')
add_para('学术前沿核心发现：', bold=True)
add_para('1. 学术界正在快速跟进：生信Agent方向已有多个团队开展工作，普瑞达有一定先发意识但不构成显著技术壁垒。')
add_para('2. RAG是关键瓶颈：学术界研究的核心共识是——RAG管线的质量决定了科研Agent的实用性。')
add_para('3. 端到端生信分析Agent仍处于早期：从AutoBA(2023)到oxo-call(2026)，学术界正在逐步攻克「LLM+生信工具链」这一课题。')
add_para('4. 标准化评估已在路上：AutoResearchBench的出现意味着这个赛道正在走向规范化，但也意味着竞争将更加激烈。')

doc.add_page_break()

# ======== 5. 产品与商业化 ========
add_heading_styled('五、产品与商业化评估', level=1)

add_heading_styled('5.1 产品竞争力分析', level=2)
add_para('普瑞达产品定位为「科研人员的AI数字员工」，涵盖文献调研、数据分析、研究辅助三大功能。')

add_styled_table(
    ['竞争要素', '评级', '说明'],
    [
        ['工作流完整性', '★★★★', '从文献调研到生信分析到研究辅助，覆盖科研完整工作流，而非单点工具'],
        ['零学习成本', '★★★★', '原生嵌入飞书/微信，用户无需安装新软件'],
        ['本地化部署', '★★★', '满足医院/药企的数据安全合规需求，但小团队维护成本高'],
        ['文献审查机制', '★★★★', 'BP核心亮点，但技术实现难度高，需验证实际效果'],
        ['差异化程度', '★★★', '在生信分析执行方面有一定差异化，但容易被同行跟进'],
    ],
    col_widths=[3, 1.5, 11]
)

add_heading_styled('5.2 商业模式评估', level=2)
add_para('三层次的商业模式设计合理：')
add_para('• 个人订阅（月付）——已跑通，验证了C端付费意愿。自媒体获客成本<10元/人是亮点。')
add_para('• 机构License（年付）——重点发力方向。机构付费=更强客户粘性+更高ARPU。机构采购周期6-12个月是主要风险。')
add_para('• 产业定制（按项目）——中期目标，面向药企/CRO。')

add_para('')
add_para('核心关注：机构采购周期长（通常6-12个月），在资金有限的早期阶段，过度依赖机构Licenses可能带来现金流压力。「个人→课题组→机构」的渗透路径合理，但转化率是关键指标。BP未披露DAU/MAU、付费转化率、客单价等核心运营指标。')

doc.add_page_break()

# ======== 6. 市场与竞争格局 ========
add_heading_styled('六、市场与竞争格局', level=1)

add_heading_styled('6.1 市场规模', level=2)
add_para('BP引用的三个市场数字具有参考价值：')
add_para('• 中国AI Agent应用市场：2025年企业级规模232亿元，CAGR 120%，2027年达655亿元')
add_para('• 中国生物信息学市场：2024年规模72.6亿元')
add_para('• 中国在学研究生409万（博士68万+硕士342万）')
add_para('但需要区分「可获得市场」（Served Addressable Market）与总市场规模。实际可触达的市场金额应聚焦于中国生命科学领域的研究型用户群体。')

add_heading_styled('6.2 竞争格局分析（增强版）', level=2)
add_para('当前AI科研助手赛道的主要竞争者可分为以下几类：')

add_styled_table(
    ['竞争者类型', '代表产品', '核心差异', '对普瑞达的威胁'],
    [
        ['国际学术搜索', 'Elicit, Scite,\nSemantic Scholar', '文献搜索+引用分析，强在论文库覆盖', '中。定位不同，偏信息检索而非分析执行'],
        ['AI科研写作', 'Paperpal, Writefull,\nJenni AI', '强在论文写作辅助', '低。侧重写作而非全链条辅助'],
        ['新型科研Agent', 'Consensus, PaperQA,\nResearchOS', 'AI驱动的科研协作/讨论', '中高。功能有重叠，但缺少原生中文IM集成'],
        ['国内竞品', '赛舵智能\n及其他未知名产品', '功能定位与普瑞达有重叠', '中。信息不完整，需进一步调研'],
        ['学术开源工作', 'AutoBA, BRAD,\noxo-call', '学术研究原型，非商业产品', '低中。可能被商业化团队快速转化'],
    ],
    col_widths=[3, 3.5, 4.5, 4.5]
)

add_heading_styled('6.3 2026年产业动态', level=2)
add_para('通过技术新闻检索，梳理了2026年5月与本项目相关的关键产业信号：')

add_para('【利好信号】')
add_para('• 国产Agent模型进入全球第一梯队：昆仑万维发布SkyClaw-v1.0，百万token上下文，深度适配Agent工作场景。意味着更低成本和更丰富的模型选择。')
add_para('• DeepSeek开发自动研究Skill：「1%是我写的，99%是Agent写的」——论文写作自动化取得突破，显示科研自动化方向的价值正在被验证。')
add_para('• AI Agent赛道整体高速增长：MIT Tech Review报道85%的组织希望3年内实现Agent化，Wired专题报道AI Agents重塑科技行业。')

add_para('【风险信号】')
add_para('• 通用Agent平台快速进化：Google I/O 2026将搜索替换为AI Agent，微软将Windows重塑为Agentic OS——底层平台都在向Agent原生演进，长期可能侵蚀垂直Agent的生存空间。')
add_para('• 企业对AI Agent的信任危机：DuckDuckGo安装量增长30%源于用户拒绝Google AI搜索；Ars Technica报道Starlette的BadHost漏洞影响数百万AI Agent。')
add_para('• 供应商锁定风险：模型API依赖度、厂商平台切换成本需考虑。')

doc.add_page_break()

# ======== 7. 团队评估 ========
add_heading_styled('七、团队评估', level=1)

add_heading_styled('7.1 核心团队胜任度', level=2)

add_styled_table(
    ['维度', '评估'],
    [
        ['创始人 朱玲玲', '中国科学技术大学博士后，干细胞命运调控方向，科研一线背景。深刻理解科研痛点，对产品方向有直觉判断。'],
        ['联合创始人 王伟光', '精通生物信息分析，与创始人多年合作，共同发表SCI论文。团队协作验证充分。'],
        ['团队背景一致性', '均为科研背景，「我们是我们用户的原型」——对科研痛点的理解是天然优势。'],
        ['关键岗位缺失', '缺少AI工程化/AI Agent系统架构方向的专职工程师——这是当前最突出的短板。'],
        ['招聘计划', '本轮融资后16个月扩至8-10人（Agent工程2人+生信研发2人+市场/产品等），节奏合理。'],
    ],
    col_widths=[3.5, 12]
)

add_heading_styled('7.2 团队评估结论', level=2)
add_para('两位创始人具备科研领域专业背景和长期合作关系，在理解用户痛点方面有天然优势。但需注意：')
add_para('• 当前2人团队缺少AI工程化经验——构建可靠、可拓展的AI Agent系统需要系统工程能力。')
add_para('• 16个月的融资后招聘计划合理，但核心工程师的招募是关键路径依赖。')
add_para('• 创始人的「科研人」身份在初期是优势，但在商业化阶段可能转为瓶颈——需要引入商业合伙人或尽快配置BD/销售能力。')

doc.add_page_break()

# ======== 8. 风险评估 ========
add_heading_styled('八、风险识别与缓解', level=1)

add_heading_styled('8.1 技术风险', level=2)
add_styled_table(
    ['风险', '等级', '说明与缓解'],
    [
        ['LLM幻觉', '高', '科研文献引用错误可能导致严重后果（如误导实验设计）。缓解：需构建严格的RAG+引用溯源管线，建立人工复核机制'],
        ['技术壁垒薄弱', '中高', '应用层Agent方案技术门槛不高，通用Agent快速进化可能侵蚀市场。缓解：深耕生信领域知识库，构建领域数据飞轮'],
        ['服务稳定性', '中', '依赖第三方模型API，API变更/中断可能影响服务。缓解：多模型切换机制，建立模型适配层'],
        ['私有化部署维护', '中', '小团队的私有化部署运维成本高。缓解：采用轻量化容器方案，建立自动化运维体系'],
    ],
    col_widths=[3, 1.5, 11.5]
)

add_heading_styled('8.2 市场与商业化风险', level=2)
add_styled_table(
    ['风险', '等级', '说明与缓解'],
    [
        ['机构采购周期', '高', '机构License转化周期长（6-12个月），早期现金流压力大。缓解：维系个人订阅收入作为现金牛，拓展课题组级别的快速决策通道'],
        ['竞品先发优势', '中', '国内外竞品已有融资和用户积累。缓解：聚焦差异化功能（文献审查机制、中文IM集成），快速迭代建立用户基础'],
        ['获客成本控制', '中', '自媒体获客成本<10元/人是亮点但规模有限。缓解：建立学术网络KOL传播体系'],
        ['市场教育成本', '中', '科研人员对AI Agent的信任建立需要时间。缓解：提供免费试用，积累口碑和案例'],
    ],
    col_widths=[3, 1.5, 11.5]
)

add_heading_styled('8.3 团队风险', level=2)
add_styled_table(
    ['风险', '等级', '说明与缓解'],
    [
        ['关键人员风险', '高', '仅2人，关键技术依赖单人。任何一位核心成员退出将严重影响项目推进。缓解：尽快扩充团队，建立知识共享机制'],
        ['商业化能力', '中高', '创始团队为纯科研背景，缺少商业运营经验。缓解：引入商业合伙人/顾问，进行商业化培训'],
        ['招聘难度', '中', 'AI工程化人才市场竞争激烈。缓解：利用中科大校友网络，提供股权激励'],
    ],
    col_widths=[3, 1.5, 11.5]
)

doc.add_page_break()

# ======== 9. 尽调建议 ========
add_heading_styled('九、下一步尽调建议', level=1)
add_para('基于初评阶段的分析，建议在下一轮尽调中重点关注以下问题：')

add_heading_styled('9.1 技术维度进一步验证', level=2)
add_para('1. 请求产品演示，重点观察文献检索结果的准确性和引用溯源能力。')
add_para('2. 确认使用的底层模型来源（GPT-4/Claude/国产模型），以及多模型之间的切换策略。')
add_para('3. 要求提供技术架构文档，重点评估Agent编排引擎的设计成熟度和可扩展性。')
add_para('4. 评估私有化部署方案的实际可行性和成本结构。')
add_para('5. 考察RAG管线的具体实现——这是决定「文献审查机制」效果的关键。')

add_heading_styled('9.2 产品与商业化验证', level=2)
add_para('1. 获取内测用户使用数据：DAU/MAU、留存率、平均使用时长、付费转化率、客单价。')
add_para('2. 访谈2-3家正在洽谈的机构客户，验证需求真实性和采购意向。')
add_para('3. 要求提供定价模型和收入预测的详细假设依据。')
add_para('4. 评估种子轮资金的续航能力（18个月runway是否足够）。')

add_heading_styled('9.3 竞争与市场验证', level=2)
add_para('1. 对BP中提到的竞品进行详细分析对比，补充竞品对比表。')
add_para('2. 调研国内已有AI科研助手产品的用户反馈，了解市场痛点满足程度。')
add_para('3. 关注2026年下半年DeepSeek等团队在科研自动化方向的商业化进展。')

add_heading_styled('9.4 推荐验证的专家方向', level=2)
add_para('1. AI Agent系统工程专家：评估技术方案的可扩展性和工程成熟度。')
add_para('2. 生物信息学研究员（非创业者）：验证产品对生信分析场景的实际覆盖能力。')
add_para('3. 科研机构IT采购决策者：了解机构采购的预算、流程和决策因素。')
add_para('4. AI创业者（已退出/有经验）：评估早期商业化和团队配置合理性。')

doc.add_page_break()

# ======== Disclaimer ========
add_para('')

# Add horizontal line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('─' * 60)
run.font.color.rgb = RGBColor(150, 150, 150)

add_para('')
add_para('免责声明', bold=True, size=Pt(11), alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_para('本报告由Deeptechnic（技研）基于公开可获取信息撰写，仅供投资决策参考。报告中的评级和判断为技术研判意见，不构成投资建议。所有判断标注了置信度，随着新信息的输入可不断修正。', size=Pt(9))
add_para('')
add_para('信息来源：项目商业计划书、arXiv学术论文检索、Semantic Scholar学术检索、技术新闻媒体检索（36氪、钛媒体、量子位、雷锋网、InfoQ、开源中国、V2EX、博客园、TechCrunch、Wired、MIT Technology Review等）、公开市场数据。', size=Pt(9))
add_para('')
add_para('报告编号：DEEPTECHNIC-20260527-PRIMDATA-AI', size=Pt(9), color=RGBColor(150, 150, 150))

# ======== SAVE ========
output_dir = '/Users/ackiles/Documents/【6-4】对接项目/安徽节目录制项目'
output_path = os.path.join(output_dir, 'AI多智能体科研助理-普瑞达PrimdataAI_技术初评报告_更新版.docx')
doc.save(output_path)
print(f'Saved to: {output_path}')
print(f'File size: {os.path.getsize(output_path)} bytes')
