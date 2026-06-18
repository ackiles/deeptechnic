# -*- coding: utf-8 -*-
"""
Deeptechnic（技研）项目初评报告生成脚本 v2.1
基于 GB/T 9704-2012 党政机关公文格式标准
版本说明：v2.1 修正了v2.0"BP综述"的问题，全部改用独立工程分析，不再转录BP宣传语
"""

from docx import Document
from docx.shared import Pt, Cm, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

LINE_SPACING = Pt(28.5)

CONFIG = {
    "docNumber": None,
    "docNumberAlign": "center",
    "title": "关于「脉冲红外连续时间视觉智能石化行业应用」项目的技术初评报告（v2.1）",
    "mainRecipients": None,
    "author": "Deeptechnic（技研）",
    "date": "2026年5月22日",
    "attachments": [],
    "contact": None,
    "contactFormat": {"bracket": True, "separator": "  "},
    "showPageNumber": True,
    "outputFile": "/Users/ackiles/Documents/AI-Space/专家/reports/2026-05-22-脉冲红外技术深研-a5cd/Deeptechnic_脉冲红外项目技术初评报告_v2.docx"
}

CONTENTS = [
    ("h1", "一、项目概要"),
    ("normal", "项目名称：脉冲红外连续时间视觉智能石化行业应用。"),
    ("normal", "本报告为工程技术可行性独立分析，不转录BP宣传语，所有判断基于公开技术原理和工程常识。"),
    ("blank", None),

    # ── 二、核心技术claim的物理可行性分析 ──
    ("h1", "二、核心技术claim的物理可行性分析"),

    ("h2", "（一）「脉冲连续摄影原理」本身是经过学术验证的原始创新。"),
    ("normal", "该原理已在可见光波段大量论文验证（Engineering 2022、CVPR、NeurIPS、TPAMI等）。核心创新是：每个像素独立进行光子积分-阈值比较-脉冲输出，用脉冲间隔倒数表示光强，而非传统相机的固定曝光时间。这一机制在理论上是正确的，已有实际芯片和相机验证。"),

    ("h2", "（二）关键claim #1：「15um像元实现与30um像元等同的NETD 20mK」——物理上值得深究。"),
    ("normal", "NETD（噪声等效温差）的物理公式简化为：NETD ∝ 1/(D*·√(A_d·Δt))，其中A_d为像元面积，Δt为积分时间。15um像元面积是30um的1/4。要维持相同NETD，理论上需要：(a) D*提升4倍（同样材料不可能），或(b) 积分时间延长4倍（与256Hz高帧率矛盾），或(c) 光学系统F数降低2倍（更大/更贵镜头）。"),
    ("normal", "BP称'通过脉冲技术特性'解决了这一矛盾，但未给出物理机理。可能的解释是：脉冲读出机制不采用传统积分-读出方式，每个像素独立触发脉冲，等效于自适应积分时间。但这需要验证——脉冲模式下NETD是否随光强（背景温度）变化？在低照度/低温差场景下NETD是否会显著劣化？这是必须向团队澄清的核心技术问题。"),

    ("h2", "（三）关键claim #2：「256Hz帧率」——冷却型焦平面的读出速度瓶颈。"),
    ("normal", "640x512 x 256Hz = 84M pixels/s的读出速率。商用冷却型MCT/T2SL焦平面的典型读出速率：640x512@30-60Hz是常规产品，100-200Hz是高端产品（如FLIR X系列科研级相机可达1000fps但采用子窗口模式）。84M pixels/s的读速在冷却红外领域属于高速，并非不可能但非常态。"),
    ("normal", "核心工程问题：脉冲读出模式与冷却型焦平面是否兼容？冷却型焦平面的暗电流比可见光CMOS大几个数量级，在77K工作温度下ROIC的噪声特性会显著变化。标准的CMOS工艺在低温下存在载流子冻析效应（carrier freeze-out），可能导致脉冲触发电路工作不稳定。这需要特种工艺（如低温CMOS或SOI），会显著增加ROIC流片成本和风险。"),

    ("h2", "（四）关键claim #3：「2-4um宽谱段覆盖800种气体」——信号竞争和干扰问题。"),
    ("normal", "红外气体成像的物理原理：气体分子吸收特定波长红外辐射，在热成像中表现为'冷云'（吸收背景辐射）。3.2-3.5um波段恰好覆盖C-H键基频伸缩振动，是绝大多数VOCs的强吸收区。扩展到2-4um（特别是2-2.5um区域）确实增加了可检测气体种类，但也带来了关键问题："),
    ("normal", "①干扰信号增加：2.0-2.5um区域包含CO2（2.0、2.7um）、H2O（2.7um）和NH3（2.2um）等非VOCs气体的强吸收线。在石化厂区环境中，这些干扰气体普遍存在，会导致大量误报。"),
    ("normal", "②气体种类'覆盖'≠可区分检测：探测器看到的是气体吸收导致的温度变化，无法分辨是哪种气体导致的吸收。BP声称可识别400+种VOCs，但实际只能检测'有气体泄漏'，不能区分具体种类。曹汛团队的光谱相机方案正是为了解决这一问题，用色散元件获取光谱维度的信息。"),
    ("normal", "③被动检测的共性问题：气体是透明还是可见取决于气体与背景的温度差。如果气体与背景温度相同（如管道泄漏后气体迅速扩散至环境温度），则完全不可见。这也是为什么OGI（光学气体成像）在检测低浓度慢泄漏时仍然很困难。"),

    ("h2", "（五）关键claim #4：「低浓度效果明显，3%甲烷在0.5L/min下可见」——实验室与现场差距。"),
    ("normal", "实验室条件下的气体泄漏检测（干净背景、控制温差的充气云室）与实际现场（高温管道、复杂背景、多气体共存、大风/雨雪天气）之间存在巨大差距。"),
    ("normal", "FLIR GFx320在实验室条件下也能检测低浓度甲烷，但现场测试中因温差不足、背景复杂而性能大幅下降。脉冲红外在实验室优于GFx320的测试结果值得关注，但缺少第三方独立测试和真实厂区现场（而非实验室）的对比数据。"),
    ("blank", None),

    # ── 三、核心技术迁移工程评估 ──
    ("h1", "三、脉冲→红外核心工程迁移难题"),
    ("normal", "BP的核心叙事是'脉冲相机原理已在可见光波段验证，现在迁移到中波红外'。这一迁移包含以下尚未验证的工程环节："),

    ("h2", "（一）ROIC重新设计。"),
    ("normal", "现有脉冲ROIC（如用于Vidar脉冲相机的读出芯片）是为可见光CMOS图像传感器设计的。迁移到红外焦平面需要：①适配MCT/T2SL探测器的暗电流特性（IR探测器暗电流比可见光高2-3个数量级）；②脉冲触发阈值需适应IR的低信号水平（IR像元电容通常更小）；③低温77K工作环境对CMOS晶体管的阈值电压、迁移率和噪声特性有显著影响；④需通过倒装焊（flip-chip）与红外探测器混合集成，引入额外的寄生电容和热应力问题。"),

    ("h2", "（二）混合集成与封装。"),
    ("normal", "脉冲ROIC芯片与红外焦平面传感器的混合集成是核心技术挑战之一。标准的IR FPA-ROIC flip-chip封装（In柱键合）对ROIC的金属层设计和pad布局有严格要求。NERC在可见光CMOS脉冲传感器的经验是否可直接迁移到IR焦平面领域，是需要工程验证的。"),

    ("h2", "（三）15um像元制冷型焦平面的供应链可用性。"),
    ("normal", "15um像元在非制冷型IR传感器（如高德红外12um）已成熟，但在制冷型MCT/T2SL中仍属前沿。国内主要制冷型供应商（11所、211所、高德、睿创、拓感等）的量产产品以30um像元为主。15um制冷型MCT/T2SL的研发状态需逐家确认。"),
    ("normal", "即使有供应商能做15um像元，还需要同时满足：2-4um宽谱段光学响应 + 兼容脉冲ROIC的读出接口 + 批量供应能力。这一组合可能构成唯一的瓶颈。"),

    ("h2", "（四）算法适配与数据瓶颈。"),
    ("normal", "脉冲视觉在可见光波段的算法栈（重建、光流、跟踪）不能直接用于红外气体检测。须重新训练针对红外气体特征的AI模型。目前仅有实验室条件下的甲烷、乙烷、丁烷数据，距离实际应用所需的百万级标注图像差距巨大。"),
    ("blank", None),

    # ── 四、供应链深度分析 ──
    ("h1", "四、供应链与成本结构分析"),

    ("h2", "（一）冷却型红外相机关键BOM。"),
    ("normal", "一套冷却型中波红外相机（640x512）的核心BOM包括：①红外焦平面探测器（MCT/T2SL FPA，含ROIC）；②斯特林制冷机（微循环式，约10-20万元/个，占系统成本~30%）；③红外光学镜头（因冷却型需冷光阑匹配，镜头设计复杂，造价~10-20万元）；④图像处理电路和接口板。"),
    ("normal", "BP定价50万元/套，这个定价在冷却型红外相机市场属于'有竞争力的国产价格'。但需注意斯特林制冷机有MTBF（平均无故障时间）限制，通常在8000-20000小时（~1-2.5年连续运行），维护更换成本约5-10万元/次。这对石化行业7x24小时连续运行场景的TCO有显著影响。"),

    ("h2", "（二）脉冲ROIC的流片成本和周期。"),
    ("normal", "定制ROIC如采用先进CMOS工艺（>180nm），全掩膜流片费用约50-100万美元/次，周期4-6个月。加上封装、测试和与FPA的混合集成验证，从流片到可用样片至少12-18个月。如果通过MPW（多项目晶圆）方式，成本可降至10-15万美元/次，但样品数量少、灵活度低。"),

    ("h2", "（三）国内供应商逐一分析。"),
    ("normal", "高德红外：国内最大红外厂商，有制冷和非制冷产线。但高德自己生产完整的红外热像仪产品，是潜在竞品而非单纯的供应商。"),
    ("normal", "睿创微纳（艾睿）：以非制冷型为主，制冷型产能较小。有MCT产线但侧重军用。"),
    ("normal", "安徽光智科技：产能最大（宣称15-20K/年），MCT+T2SL双路线。但产能数据含材料、衬底、传感器等叠加统计，实际FPA出货量可能低一个数量级。"),
    ("normal", "拓感科技：二类超晶格路线，适合本项目第二阶段方案。产能较小（2-3K/年）。"),
    ("normal", "整体判断：国产IR FPA供应充足，但'15um + 宽谱段 + 脉冲ROIC接口'的组合需求可能需定制开发，不再是标准货架产品。"),

    ("h2", "（四）关键器件自主风险。"),
    ("normal", "斯特林制冷机：国产替代已完成（如昆明物理所、安徽万瑞等），但高可靠性长寿命制冷机仍以进口为主（Ricor、Thales、Leybold等，受出口管制影响）。"),
    ("normal", "红外光学材料：锗（Ge）、硫化锌（ZnS）、硒化锌（ZnSe）国内供应充足。"),
    ("normal", "冷光阑：需精密加工，国内可做。"),
    ("normal", "整体来看，没有'被卡脖子'的供应链关键点，但脉冲ROIC的定制化成本和时间是主要风险。"),
    ("blank", None),

    # ── 五、团队工程能力评估 ──
    ("h1", "五、团队工程能力独立评估"),

    ("h2", "（一）原创IP持有者投入不足。"),
    ("normal", "黄铁军教授（首席科学家，兼职10%）是本项目核心IP的唯一源头。他的兼职投入意味着：(a) 关键技术决策依赖碎片化时间；(b) 知识产权转让/独占授权安排可能尚未完成法律程序；(c) 如果NERC或其他团队用同一技术做类似项目（如可见光领域的机器视觉），可能存在IP冲突。"),

    ("h2", "（二）红外硬件工程能力存在结构性缺口。"),
    ("normal", "逐一评估全职核心成员的硬件工程背景："),
    ("normal", "杨圣伟（CEO，全职）：大华营销/管理背景，负责产品规划、招投标、客户关系，不具备红外焦平面或光学系统的技术决策能力。"),
    ("normal", "张鹏飞（CTO，全职）：北邮硕士，脉冲信号处理方向，偏软件/算法，无红外探测器或硬件工程经验。"),
    ("normal", "姜铁刚（首席架构师，全职）：浙大学士，大华产品线总监背景，偏软件产品管理和营销。"),
    ("normal", "团队无一名全职成员有红外焦平面探测器设计、ROIC设计、混合集成、斯特林制冷机系统等领域的实战经验。"),
    ("normal", "这意味着项目的核心技术实现（红外焦平面设计与ROIC定制）只能依赖NERC的外部协作，NERC承担了实际上的'工程实现'角色但并非项目公司成员。"),

    ("h2", "（三）与竞品团队的差距。"),
    ("normal", "对比曹汛团队（PMVIS光谱相机）：曹汛本人为南京大学教授全职创业，技术带头人全时投入。同时曹汛团队已完成公司注册和融资，有明确的商业化推进。本项目的核心人员投入度和公司化进度明显落后。"),
    ("blank", None),

    # ── 六、竞争格局的独立分析 ──
    ("h1", "六、竞争格局独立分析"),

    ("h2", "（一）OLI（Optical Gas Imaging）竞争格局。"),
    ("normal", "全球OGI市场主要参与者：Teledyne FLIR（美国，占据全球最大份额）、Opgal（以色列，EyeCheck系列）、Honeywell（美国，Rebellion系列）、飒特（中国，EX系列）、紫来（中国，M系列）、大立科技（中国，DM系列）。"),
    ("normal", "FLIR GFx320虽贵（~100万元），但作为行业标杆产品已有大规模部署（全球数万套），算法和软件生态成熟。国产竞品（飒特EX350、紫来M320）价格已降至~50万元，在石化招标中频繁中标，初步完成国产替代。"),

    ("h2", "（二）曹汛PMVIS光谱方案vs脉冲红外方案的实质对比。"),
    ("normal", "BP'竞争分析'章节仅对比了4款产品（脉冲红外、FLIR、飒特、紫来），完全未提及曹汛团队。这一缺口是严重的，因为曹汛方案在气体检测领域与脉冲红外方案存在直接替代关系："),
    ("normal", "脉冲红外优势：基于气体吸收导致的'温度变化'成像，可提供高帧率动态图像，适合定位泄漏源和追踪扩散路径。"),
    ("normal", "光谱相机优势：通过光谱分解可区分不同气体种类，减少误报；可实现浓度定量估算；更多维度的数据有利于符合法规和ESG报告需求。"),
    ("normal", "本质区别：脉冲红外方案提供'哪里漏了+漏到什么程度'的信息，光谱方案提供'漏的是什么+浓度多少'的信息。在石化客户的实际场景中，两者不是二选一，而是都可以卖——但客户预算有限，形成竞争关系。"),

    ("h2", "（三）脉冲红外的核心差异化优势评估。"),
    ("normal", "高帧率（256Hz）实现'动中测'是当前所有OGI产品的空缺，确实存在差异化。但需验证：石化行业客户是否有显著的'动中测'需求？在线式防爆云台本身就有运动能力（水平和垂直转动），常规15-30Hz帧率是否已够用？"),
    ("normal", "'动中测'的主要价值场景是车载/无人机载平台——但这些平台的部署数量远少于固定安装点。手提式OGI是当前最大量的需求（单人巡检），对动中测无需求。"),
    ("blank", None),

    # ── 七、综合评价与建议 ──
    ("h1", "七、综合评价与建议"),

    ("h2", "（一）综合评价。"),
    ("normal", "技术先进性4.5/5：脉冲连续摄影原理是真正的原始创新，但在红外迁移的几个关键工程环节（ROIC低温工作、15um NETD保持、宽谱段抗干扰）物理可行性尚未验证。"),
    ("normal", "工程化可行性2.5/5：从可见光脉冲原理验证到红外商用产品之间，存在大量未跨过的工程鸿沟。ROIC定制的成本和周期、15um制冷型焦平面的货架可用性、长周期运行可靠性都是高风险。"),
    ("normal", "供应链3.5/5：国产红外供应链整体替代完成，但脉冲ROIC定制化带来的成本和时间风险不可忽略。"),
    ("normal", "市场机会3.5/5：OGI市场政策驱动明确（LDAR新标准），真实需求存在。但已有多个国产成熟竞品（飒特、紫来、大立），曹汛团队也已成立公司。脉冲高帧率差异化定位的市场溢价能力需验证。"),
    ("normal", "团队执行能力2.5/5：原创IP持有者和关键算法负责人都兼职投入，全职团队在中波红外硬件工程的关键岗位上存在结构性能力缺口。CEO的营销管理背景与核心技术型创业项目的需求不够匹配。"),
    ("normal", "综合评分：3.0/5分（满分5分）。"),

    ("h2", "（二）关键待验证问题（按优先级排序）。"),
    ("normal", "P0级（必须回答，否则不可推进）："),
    ("normal", "① 15um像元MCT/T2SL焦平面在NETD 20mK下256Hz读出的物理可行性——请团队给出理论上限的工程计算，而非宣称性结论。"),
    ("normal", "② 脉冲ROIC迁移到低温（77K）工作环境的工程方案——是否采用特种CMOS工艺？低温下脉冲触发电路的工作特性变化？"),
    ("normal", "③ 黄铁军教授的脉冲IP是否有正式转让/独占授权给项目公司的法律协议——NERC是否保留将该IP用于其他商业项目的权利？"),

    ("h2", "（三）建议。"),
    ("normal", "1. 建议列为「有条件跟进」项目。"),
    ("normal", "2. 安排红外探测器领域独立专家（如211所或高德红外的退休技术负责人）对ROIC迁移方案进行技术评审。"),
    ("normal", "3. 要求团队提供一份工程可行性分析报告，包含：NETD-帧率-像元尺寸相互制约的理论计算、ROIC低温特性仿真数据、15um焦平面供应商的正式报价和交期。"),
    ("normal", "4. 验证NERC的合作模式：NERC在项目中是'协作开发方'还是'供应商'？ROIC的IP归属？流片费用谁承担？"),
    ("normal", "5. 要求团队补充对曹汛光谱相机方案的正式竞争分析，说明差异化定位策略。"),
    ("normal", "6. 如推进至下一轮，建议安排一次样机实地演示+独立第三方性能测试（与FLIR GFx320和飒特EX350的盲测对比）。"),
    ("normal", "7. 建议团队引入全职红外硬件工程方向的核心成员（至少1名有制冷型焦平面/ROIC开发经验的资深工程师），这应是推进尽调的前置条件。"),

    ("h2", "（四）初评结论。"),
    ("normal", "脉冲连续摄影原理有真实的技术创新价值，在国内红外气体检测这个政策驱动的增长赛道中有差异化定位的市场空间。但当前团队配置、工程化进度和核心技术的工程可行性验证均存在显著风险。建议在P0级工程问题得到明确回答、团队补齐硬件工程能力缺口后，再决定是否推进至深入的专家访谈话第二轮。"),
]

# ── 辅助函数 ──
LEFT_DOUBLE = "\u201C"
RIGHT_DOUBLE = "\u201D"
LEFT_SINGLE = "\u2018"
RIGHT_SINGLE = "\u2019"

def normalize_chinese_quotes(text):
    if not text:
        return text
    result = []
    double_count = 0
    single_count = 0
    i = 0
    while i < len(text):
        c = text[i]
        if c == '"':
            result.append(LEFT_DOUBLE if double_count % 2 == 0 else RIGHT_DOUBLE)
            double_count += 1
        elif c == "'":
            result.append(LEFT_SINGLE if single_count % 2 == 0 else RIGHT_SINGLE)
            single_count += 1
        else:
            result.append(c)
        i += 1
    return ''.join(result)

def set_run_font(run, font_name="仿宋_GB2312", font_size=16, bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)

def create_paragraph(doc, text, font_name="仿宋_GB2312", font_size=16,
                     alignment=WD_ALIGN_PARAGRAPH.LEFT,
                     first_line_indent=Twips(640), bold=False,
                     space_before=0, space_after=0, right_indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = alignment
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent and first_line_indent > 0:
        p.paragraph_format.first_line_indent = first_line_indent
    else:
        p.paragraph_format.first_line_indent = Cm(0)
    if right_indent is not None:
        p.paragraph_format.right_indent = right_indent
    if text:
        text = normalize_chinese_quotes(text)
        run = p.add_run(text)
        set_run_font(run, font_name, font_size, bold)
    return p

def create_blank_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p

def split_h2_title_and_body(text):
    if not text:
        return text, None
    if '\n' in text:
        parts = text.split('\n', 1)
        return parts[0].strip(), parts[1].strip()
    for i, c in enumerate(text):
        if c == '。':
            remaining = text[i + 1:].strip()
            if remaining:
                return text[:i + 1], remaining
            else:
                return text, None
    return text, None

def add_h2_with_body(doc, text):
    title_text, body_text = split_h2_title_and_body(text)
    title_text = normalize_chinese_quotes(title_text)
    if body_text is not None:
        body_text = normalize_chinese_quotes(body_text)
    if body_text is None:
        return create_paragraph(doc, title_text, font_name="楷体_GB2312", font_size=16, first_line_indent=Twips(640), bold=False)
    if '\n' in text:
        create_paragraph(doc, title_text, font_name="楷体_GB2312", font_size=16, first_line_indent=Twips(640), bold=False)
        return create_paragraph(doc, body_text, font_name="仿宋_GB2312", font_size=16, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Twips(640), bold=False)
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Twips(640)
    run_title = p.add_run(title_text)
    set_run_font(run_title, "楷体_GB2312", 16, bold=False)
    run_body = p.add_run(body_text)
    set_run_font(run_body, "仿宋_GB2312", 16, bold=False)
    return p

def add_content(doc, item):
    content_type, text = item
    if content_type == "blank":
        return create_blank_line(doc)
    elif content_type == "h1":
        return create_paragraph(doc, text, font_name="黑体", font_size=16, first_line_indent=Twips(640), bold=False)
    elif content_type == "h2":
        return add_h2_with_body(doc, text)
    elif content_type == "normal":
        return create_paragraph(doc, text, font_name="仿宋_GB2312", font_size=16, first_line_indent=Twips(640), bold=False)
    return create_paragraph(doc, text)

def set_page_margins(doc, top=Cm(3.7), bottom=Cm(3.5), left=Cm(2.8), right=Cm(2.6), header_distance=Cm(1.5), footer_distance=Cm(2.5)):
    for section in doc.sections:
        section.top_margin = top
        section.bottom_margin = bottom
        section.left_margin = left
        section.right_margin = right
        section.header_distance = header_distance
        section.footer_distance = footer_distance

def add_page_number(doc):
    section = doc.sections[0]
    sectPr = section._sectPr
    titlePg = OxmlElement('w:titlePg')
    sectPr.append(titlePg)
    evenAndOdd = OxmlElement('w:evenAndOddHeaders')
    sectPr.append(evenAndOdd)

    def make_page_number_fldChar(run_elem, fld_type="current"):
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE \\* ARABIC ' if fld_type == "current" else ' NUMPAGES \\* ARABIC '
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        run_elem.append(fldChar_begin)
        run_elem.append(instrText)
        run_elem.append(fldChar_end)

    def set_footer_run_font(run_elem, size_pt=14):
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), '宋体')
        rFonts.set(qn('w:hAnsi'), '宋体')
        rFonts.set(qn('w:eastAsia'), '宋体')
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(size_pt * 2))
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(size_pt * 2))
        rPr.append(rFonts)
        rPr.append(sz)
        rPr.append(szCs)
        run_elem.append(rPr)

    def build_footer_paragraph(footer, alignment):
        for p in footer.paragraphs:
            p._element.getparent().remove(p._element)
        pPr_elem = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), alignment)
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:line'), '240')
        spacing.set(qn('w:lineRule'), 'auto')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:firstLine'), '0')
        ind.set(qn('w:left'), '0')
        pPr.append(jc)
        pPr.append(spacing)
        pPr.append(ind)
        pPr_elem.append(pPr)
        footer._element.append(pPr_elem)
        return pPr_elem

    ONE_IDEOGRAPH_SPACE = '\u3000'
    odd_footer = section.footer
    p_odd = build_footer_paragraph(odd_footer, 'right')
    r1 = OxmlElement('w:r')
    set_footer_run_font(r1)
    t1 = OxmlElement('w:t')
    t1.set(qn('xml:space'), 'preserve')
    t1.text = '\u2014 '
    r1.append(t1)
    p_odd.append(r1)
    r2 = OxmlElement('w:r')
    set_footer_run_font(r2)
    make_page_number_fldChar(r2)
    p_odd.append(r2)
    r3 = OxmlElement('w:r')
    set_footer_run_font(r3)
    t3 = OxmlElement('w:t')
    t3.set(qn('xml:space'), 'preserve')
    t3.text = ' \u2014' + ONE_IDEOGRAPH_SPACE
    r3.append(t3)
    p_odd.append(r3)

    even_footer = section.even_page_footer
    if even_footer:
        p_even = build_footer_paragraph(even_footer, 'left')

def generate_document():
    print("=== Deeptechnic 项目技术初评报告生成 v2.1 ===")
    print("标题:", CONFIG["title"])
    print("输出:", CONFIG["outputFile"])
    print("========================================\n")
    doc = Document()
    set_page_margins(doc)
    create_paragraph(doc, "Deeptechnic（技研）", font_name="黑体", font_size=22, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0, bold=False)
    create_blank_line(doc)
    if CONFIG["docNumber"]:
        create_paragraph(doc, CONFIG["docNumber"], alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0)
    create_blank_line(doc)
    create_blank_line(doc)
    create_paragraph(doc, CONFIG["title"], font_name="方正小标宋简体", font_size=22, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0, bold=False)
    create_blank_line(doc)
    for item in CONTENTS:
        add_content(doc, item)
    SIGN_RIGHT_INDENT = Twips(1280)
    create_blank_line(doc)
    create_blank_line(doc)
    create_blank_line(doc)
    if CONFIG["author"]:
        create_paragraph(doc, CONFIG["author"], alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_line_indent=0, right_indent=SIGN_RIGHT_INDENT)
    if CONFIG["date"]:
        create_paragraph(doc, CONFIG["date"], alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_line_indent=0, right_indent=SIGN_RIGHT_INDENT)
    if CONFIG.get("showPageNumber", True):
        add_page_number(doc)
    doc.save(CONFIG["outputFile"])
    print("生成成功:", CONFIG["outputFile"])

if __name__ == "__main__":
    generate_document()
